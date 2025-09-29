"""DOCX parser that extracts text, headings, and metadata."""

from __future__ import annotations

import datetime as _dt
from pathlib import Path
from typing import Any

from . import DocumentSection, PageContent, ParsedDocument, ParserError


def _extract_core_properties(doc: Any) -> dict[str, Any]:
    properties = doc.core_properties
    metadata: dict[str, Any] = {}
    for attr in (
        "title",
        "subject",
        "author",
        "category",
        "comments",
        "keywords",
        "last_modified_by",
    ):
        value = getattr(properties, attr, None)
        if value:
            metadata[attr] = value

    for attr in ("created", "modified", "last_printed"):
        value = getattr(properties, attr, None)
        if isinstance(value, _dt.datetime):
            if value.tzinfo is None:
                value = value.replace(tzinfo=_dt.timezone.utc)
            metadata[attr] = value.astimezone(_dt.timezone.utc).isoformat()
    return metadata


def parse_docx(path: Path) -> ParsedDocument:
    """Parse a Microsoft Word document and extract headings and body content."""

    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - defensive
        raise ParserError("python-docx is required to parse DOCX files") from exc

    document = Document(path)
    metadata = _extract_core_properties(document)

    sections: list[DocumentSection] = []
    current_section: DocumentSection | None = None
    body_parts: list[str] = []
    char_offset = 0
    line_number = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.rstrip()
        if not text:
            char_offset += 1
            line_number += 1
            continue
        body_parts.append(text)
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.lower().startswith("heading"):
            try:
                level = int("".join(filter(str.isdigit, style_name)))
            except ValueError:
                level = None
            current_section = DocumentSection(
                title=text,
                content="",
                level=level,
                page_number=1,
                start_offset=char_offset,
                end_offset=char_offset + len(text),
                line_start=line_number + 1,
                line_end=line_number + 1,
            )
            sections.append(current_section)
            char_offset += len(text) + 1
            line_number += 1
            continue
        if current_section is None:
            current_section = DocumentSection(
                title=None,
                content="",
                level=None,
                page_number=1,
                start_offset=char_offset,
                line_start=line_number + 1,
            )
            sections.append(current_section)
        if current_section.content:
            current_section.content += "\n"
        current_section.content += text
        current_section.end_offset = char_offset + len(text)
        current_section.line_end = line_number + 1
        if current_section.line_start is None:
            current_section.line_start = line_number + 1
        char_offset += len(text) + 1
        line_number += 1

    # Remove empty trailing sections to keep output compact.
    sections = [section for section in sections if section.content or section.title]
    pages = [PageContent(number=1, text="\n".join(body_parts))]
    combined = "\n".join(body_parts)
    metadata.setdefault("paragraph_count", len(document.paragraphs))
    return ParsedDocument(text=combined, metadata=metadata, sections=sections, pages=pages)
